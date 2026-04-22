from .models import SchoolMemory, SchoolSettings
import logging
import os
import random
import time
import hashlib
from django.core.cache import cache
from PyPDF2 import PdfReader
from docx import Document
from bs4 import BeautifulSoup

try:
    import requests
    HAS_OPENROUTER = True
except ImportError:
    HAS_OPENROUTER = False

logger = logging.getLogger(__name__)

# OpenRouter: نموذج DeepSeek عبر الاشتراك
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_MODEL = "deepseek/deepseek-chat"


class AIService:
    """
    خدمة الذكاء الاصطناعي عبر OpenRouter فقط (نموذج DeepSeek).
    يستخدم المفتاح OPENROUTER_API_KEY من متغيرات البيئة أو ملف .env.
    """

    def __init__(self, user=None):
        self.user = user
        self.settings = SchoolSettings.objects.first() if SchoolSettings.objects.exists() else None
        self.openrouter_keys = self._load_keys("OPENROUTER_API_KEY")
        # مفتاح الإدارة (Management Key) لقراءة الرصيد من /credits
        self.openrouter_mgmt_keys = self._load_keys("OPENROUTER_MANAGEMENT_KEY")

    def _load_keys(self, prefix):
        keys = []
        k1 = os.environ.get(prefix)
        if k1 and k1.strip():
            keys.append(k1.strip().strip("'").strip('"'))
        i = 2
        while True:
            k = os.environ.get(f"{prefix}_{i}")
            if not k or not k.strip():
                break
            keys.append(k.strip().strip("'").strip('"'))
            i += 1
        return keys

    def get_openrouter_balance(self):
        """رصيد OpenRouter (المتبقي من الاعتماد). واجهة الصحيح: GET /api/v1/key"""
        if not self.openrouter_keys or not HAS_OPENROUTER:
            return None
        cache_key = "openrouter_balance"
        cached = cache.get(cache_key)
        if cached is not None:
            return cached
        key = self.openrouter_keys[0]
        try:
            resp = requests.get(
                "https://openrouter.ai/api/v1/key",
                headers={"Authorization": f"Bearer {key}"},
                timeout=10,
            )
            if resp.status_code != 200:
                logger.warning("OpenRouter key endpoint returned %s: %s", resp.status_code, resp.text[:300])
                return None
            data = resp.json()
            # الاستجابة تحتوي على data: { limit_remaining, limit, usage, ... }
            info = data.get("data") if isinstance(data.get("data"), dict) else data
            if not info:
                return None
            limit_remaining = info.get("limit_remaining")
            if limit_remaining is not None:
                try:
                    balance = round(float(limit_remaining), 4)
                    cache.set(cache_key, balance, timeout=300)
                    return balance
                except (TypeError, ValueError):
                    pass
            limit = info.get("limit")
            usage = info.get("usage")
            if limit is not None and usage is not None:
                try:
                    balance = round(float(limit) - float(usage), 4)
                    cache.set(cache_key, balance, timeout=300)
                    return balance
                except (TypeError, ValueError):
                    pass
        except requests.RequestException as e:
            logger.warning("OpenRouter balance request failed: %s", e)
        except Exception as e:
            logger.exception("OpenRouter balance error: %s", e)
        return None

    def get_openrouter_balance_info(self):
        """
        معلومات الرصيد مع تشخيص واضح لواجهة لوحة القيادة.
        Returns: { ok: bool, balance: float|None, message: str, key_present: bool }
        """
        key_present = bool(self.openrouter_keys and self.openrouter_keys[0])
        mgmt_present = bool(self.openrouter_mgmt_keys and self.openrouter_mgmt_keys[0])
        if not HAS_OPENROUTER:
            return {'ok': False, 'balance': None, 'message': 'مكتبة requests غير متوفرة في الخادم.', 'key_present': key_present}
        if not key_present and not mgmt_present:
            return {'ok': False, 'balance': None, 'message': 'لم يتم ضبط مفاتيح OpenRouter في ملف .env (OPENROUTER_API_KEY أو OPENROUTER_MANAGEMENT_KEY).', 'key_present': False}

        cache_key_info = "openrouter_balance_info"
        cache_key_balance = "openrouter_balance"
        cached = cache.get(cache_key_info)
        if cached and isinstance(cached, dict) and 'ok' in cached:
            return cached

        # (1) الأفضل: /api/v1/credits يحتاج Management Key (يرجع total_credits و total_usage)
        if mgmt_present:
            mgmt_key = self.openrouter_mgmt_keys[0]
            try:
                resp = requests.get(
                    "https://openrouter.ai/api/v1/credits",
                    headers={
                        "Authorization": f"Bearer {mgmt_key}",
                        "Accept": "application/json",
                        "X-Title": "baza-app",
                        "User-Agent": "baza-app/1.0",
                    },
                    timeout=12,
                )
                if resp.status_code == 200:
                    data = resp.json()
                    info_obj = data.get("data") if isinstance(data.get("data"), dict) else data
                    total_credits = info_obj.get("total_credits")
                    total_usage = info_obj.get("total_usage")
                    if total_credits is not None and total_usage is not None:
                        balance = round(float(total_credits) - float(total_usage), 4)
                        cache.set(cache_key_balance, balance, timeout=300)
                        info = {'ok': True, 'balance': balance, 'message': 'تم تحديث الرصيد (credits).', 'key_present': True, 'stale': False}
                        cache.set(cache_key_info, info, timeout=180)
                        return info
                elif resp.status_code in (401, 403):
                    # سنكمل بالـ API key endpoint كخطة بديلة
                    pass
                else:
                    # 5xx: حاول عرض آخر رصيد معروف
                    if resp.status_code >= 500:
                        last_bal = cache.get(cache_key_balance)
                        if last_bal is not None:
                            info = {'ok': True, 'balance': last_bal, 'message': f"تعذر جلب الرصيد الآن (credits HTTP {resp.status_code}). تم عرض آخر رصيد معروف.", 'key_present': True, 'stale': True}
                            cache.set(cache_key_info, info, timeout=60)
                            return info
            except requests.RequestException:
                # نكمل للخطة البديلة
                pass

        # (2) بديل: /api/v1/key (قد لا يعكس الرصيد الحقيقي أو قد يتعطل)
        if not key_present:
            return {'ok': False, 'balance': None, 'message': 'لإظهار الرصيد بدقة، أنشئ Management Key في OpenRouter وضعه في OPENROUTER_MANAGEMENT_KEY داخل .env.', 'key_present': False}

        key = self.openrouter_keys[0]
        try:
            resp = requests.get(
                "https://openrouter.ai/api/v1/key",
                headers={
                    "Authorization": f"Bearer {key}",
                    "Accept": "application/json",
                    "X-Title": "baza-app",
                    "User-Agent": "baza-app/1.0",
                },
                timeout=12,
            )
            if resp.status_code != 200:
                # إن كان الخلل من جهة OpenRouter (5xx) نحاول إرجاع آخر رصيد معروف من الكاش بدل إظهار N/A
                if resp.status_code >= 500:
                    last_bal = cache.get(cache_key_balance)
                    if last_bal is not None:
                        info = {
                            'ok': True,
                            'balance': last_bal,
                            'message': f"تعذر جلب الرصيد الآن (OpenRouter HTTP {resp.status_code}). تم عرض آخر رصيد معروف.",
                            'key_present': True,
                            'stale': True,
                        }
                        cache.set(cache_key_info, info, timeout=60)
                        return info
                msg = f"فشل جلب الرصيد من OpenRouter (HTTP {resp.status_code})."
                if resp.status_code in (401, 403):
                    msg += " المفتاح غير صالح أو لا يملك صلاحية."
                if not mgmt_present:
                    msg += " ملاحظة: لعرض الرصيد بدقة وثبات، استعمل OPENROUTER_MANAGEMENT_KEY مع /credits."
                info = {'ok': False, 'balance': None, 'message': msg, 'key_present': True}
                cache.set(cache_key_info, info, timeout=120)
                return info

            data = resp.json()
            info_obj = data.get("data") if isinstance(data.get("data"), dict) else data
            limit_remaining = info_obj.get("limit_remaining")
            if limit_remaining is None:
                limit = info_obj.get("limit")
                usage = info_obj.get("usage")
                if limit is not None and usage is not None:
                    limit_remaining = float(limit) - float(usage)
            balance = round(float(limit_remaining), 4) if limit_remaining is not None else None
            info = {'ok': balance is not None, 'balance': balance, 'message': 'تم تحديث الرصيد.', 'key_present': True, 'stale': False}
            cache.set(cache_key_info, info, timeout=180)
            return info
        except requests.RequestException as e:
            info = {'ok': False, 'balance': None, 'message': f"تعذر الاتصال بـ OpenRouter: {e}", 'key_present': True}
            cache.set(cache_key_info, info, timeout=120)
            return info
        except Exception:
            info = {'ok': False, 'balance': None, 'message': "حدث خطأ داخلي أثناء جلب الرصيد.", 'key_present': True}
            cache.set(cache_key_info, info, timeout=120)
            return info

    def get_rag_context(self, query):
        keywords = query.split()
        matches = SchoolMemory.objects.filter(title__icontains=keywords[0])
        context_text = "\n".join([f"[{m.category}] {m.title}: {m.content}" for m in matches[:3]])
        return context_text

    def generate_response(self, system_instruction, user_query, rag_enabled=True, mode=None, page_context=None):
        if mode == 'bot_helper':
            enhanced_instruction = "أنت مساعد ذكي مدمج في واجهة تطبيق تسيير مدرسة جزائرية (مسار). هدفك هو إرشاد المستخدم ومساعدته في فهم الشاشة الحالية أو الإجابة على أسئلته باختصار ووضوح."
            if page_context:
                enhanced_instruction += f"\n\nالمستخدم يتواجد حالياً في الصفحة/الرابط التالي: {page_context}\nقدم إجابتك بناءً على هذا السياق إذا كان السؤال غير واضح."
            system_instruction = enhanced_instruction
            rag_enabled = False

        context = ""
        if rag_enabled:
            context = self.get_rag_context(user_query)

        directives = "- Concise for greetings. Detailed for plans."
        prompt = f"""
        Role: School Director Consultant.
        Context: {system_instruction}
        Data: {context}
        Query: {user_query}
        Directives: {directives}
        """

        prompt_hash = hashlib.sha256(prompt.encode('utf-8')).hexdigest()
        cache_key = f"ai_response_{prompt_hash}"
        cached_response = cache.get(cache_key)
        if cached_response:
            return cached_response

        final_response = None
        insufficient_funds = False
        last_error = None

        if self.openrouter_keys and HAS_OPENROUTER:
            final_response, insufficient_funds, last_error = self._try_openrouter(prompt)

        if final_response:
            cache.set(cache_key, final_response, timeout=172800)
            return final_response

        if insufficient_funds:
            return "INSUFFICIENT_FUNDS_ERROR"
        if last_error:
            return f"⚠️ عذراً، فشل الاتصال بالذكاء الاصطناعي (OpenRouter).\nالسبب: {last_error}"
        return "⚠️ عذراً، لم أتمكن من الاتصال بـ OpenRouter. يرجى التأكد من صحة المفتاح OPENROUTER_API_KEY في ملف .env ومن اتصال الإنترنت."

    def _handle_bot_helper_local(self, query):
        import re
        q = query.lower()

        if re.search(r'(إضافة|جديد|تسجيل)\s*(تلميذ|طالب|متعلم)', q):
            return "لإضافة تلميذ جديد: اذهب إلى 'تسيير التلاميذ' من القائمة الجانبية، انزل إلى أسفل نموذج الإدخال واضغط على زر 'جديد' (الأخضر)، ثم املأ البيانات واضغط 'حفظ'."

        if re.search(r'(حذف|ازالة|مسح)\s*(تلميذ|طالب|متعلم)', q):
            return "لحذف تلميذ: في واجهة 'تسيير التلاميذ'، حدد المربع بجانب اسم التلميذ في الجدول، ثم اضغط على زر 'حذف' (الأحمر) في شريط الإجراءات بالأعلى."

        if re.search(r'(تعديل|تغيير)\s*(تلميذ|طالب|متعلم)', q):
            return "لتعديل بيانات تلميذ: في 'تسيير التلاميذ'، اضغط على صف التلميذ في الجدول، سيظهر لك زر 'تعديل' (برتقالي)، اضغط عليه، عدّل البيانات، ثم اضغط 'حفظ'."

        if re.search(r'(صورة|صور)\s*(تلميذ|طالب|متعلم)', q):
            return "لإضافة صورة لتلميذ: أثناء إضافته أو تعديله، ستجد أزراراً تحت إطار الصورة في الجهة اليسرى (أو اليمنى حسب الاتجاه)، يمكنك تحميل صورة من جهازك أو التقاطها عبر الكاميرا."

        if re.search(r'(طباعة|بطاقة|بطاقات)', q):
            return "لطباعة بطاقات التلاميذ: من واجهة 'تسيير التلاميذ'، حدد التلاميذ المعنيين من الجدول، ثم اضغط على زر 'طباعة البطاقات' (البنفسجي) في الأعلى."

        if re.search(r'(اضافة|تسجيل|جديد)\s*(موظف|استاذ|أستاذ|عامل|اداري)', q):
            return "لإضافة موظف: اذهب إلى 'الموارد البشرية' من القائمة الجانبية، ثم اضغط على زر 'إضافة موظف' واكتب بياناته يدوياً."

        if re.search(r'(اسناد|أقسام|تدريس)\s*(موظف|استاذ|أستاذ)', q):
            return "لإسناد أقسام لأستاذ: في واجهة 'الموارد البشرية'، ستجد زر 'تعديل الإسناد' الذي يظهر لك جدولاً بكل الأساتذة مع مربعات نصية، أو يمكنك النقر على أيقونة السبورة بجانب اسم الأستاذ."

        if re.search(r'(استيراد|اكسل|excel|قوائم)', q):
            return "لاستيراد قوائم (تلاميذ أو أساتذة أو إسناد): يوجد أزرار استيراد خاصة في أعلى صفحة 'تسيير التلاميذ' أو 'الموارد البشرية'. تأكد من أن ملفك بصيغة Excel (أو PDF/Word للإسناد الشامل)."

        if re.search(r'(مطعم|كانتين|نصف داخلي)', q):
            return "لتسجيل حضور التلاميذ في المطعم: توجه إلى واجهة 'المطعم المدرسي'. تذكر أن التلميذ يجب أن تكون صفته 'نصف داخلي' ليظهر هناك."

        if re.search(r'(مكتبة|اعارة|إعارة|كتاب)', q):
            return "لإعارة الكتب: اذهب إلى واجهة 'المكتبة'. ستحتاج لإدخال رقم تعريف التلميذ لتسجيل الإعارة أو الإرجاع."

        if re.search(r'(تحديث|تحديثات|موافقة)', q):
            return "صفحة 'التحديثات المعلقة' (Pending Updates): تظهر للمدير فقط في القائمة الجانبية، وهي للموافقة على التعديلات (مثل إضافة/حذف تلميذ) التي قام بها المستخدمون الآخرون."

        if re.search(r'(مهام|مهمة)', q):
            return "لإسناد المهام: استخدم واجهة 'المهام' من القائمة الجانبية لإرسال مهام للطاقم ومتابعة تنفيذها."

        if re.search(r'(سلام|مرحبا|أهلا|مساعدة)', q):
            return "أهلاً بك! أنا المساعد التقني المحلي (بدون إنترنت) لتطبيق 'تسيير متوسطة بوشنافة عمر'. يمكنني إرشادك لأماكن الأزرار وكيفية عمل الواجهات. اسألني مثلاً: 'كيف أحذف تلميذ؟' أو 'أين أجد الإسناد؟'."

        return "عذراً، لم أفهم سؤالك. يرجى سؤالي عن وظيفة محددة في التطبيق مثل: 'كيف أضيف تلميذ؟'، 'كيف أعدل الإسناد؟'، أو 'كيف أطبع البطاقات؟'."

    def _try_openrouter(self, prompt):
        """استدعاء OpenRouter (نموذج DeepSeek). يُرجع (response_text, insufficient_funds, last_error)."""
        keys = list(self.openrouter_keys)
        random.shuffle(keys)
        last_error = None

        for key in keys:
            try:
                resp = requests.post(
                    OPENROUTER_URL,
                    headers={
                        "Authorization": f"Bearer {key}",
                        "Content-Type": "application/json",
                        "HTTP-Referer": "http://localhost",
                        "X-Title": "School Management Agent",
                    },
                    json={
                        "model": OPENROUTER_MODEL,
                        "messages": [{"role": "user", "content": prompt}],
                    },
                    timeout=60,
                )
                if resp.status_code == 200:
                    data = resp.json()
                    if "choices" in data and len(data["choices"]) > 0:
                        return data["choices"][0]["message"]["content"], False, None
                else:
                    last_error = f"HTTP {resp.status_code}: {resp.text[:200]}"
                    if resp.status_code == 402 or "insufficient_quota" in resp.text.lower() or "balance" in resp.text.lower():
                        return None, True, last_error
                    logger.warning("OpenRouter API error: %s", last_error)
            except Exception as e:
                last_error = str(e)
                logger.warning("OpenRouter request exception: %s", e)
                continue

        return None, False, last_error


def analyze_assignment_document(a):
    pass


def analyze_global_assignment(f):
    pass


def analyze_global_assignment_content(file_path):
    from students.utils_tools.smart_assignment_analyzer import extract_from_excel, extract_from_word, extract_from_pdf
    ext = os.path.splitext(file_path)[1].lower()

    raw_text = ""
    if ext in ['.xlsx', '.xls']:
        return extract_from_excel(file_path)
    elif ext == '.docx':
        from docx import Document
        try:
            doc = Document(file_path)
            raw_text = "\n".join([p.text for p in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    raw_text += " | ".join([cell.text for cell in row.cells]) + "\n"
        except Exception:
            pass
    elif ext == '.pdf':
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    raw_text += (page.extract_text() or "") + "\n"
        except Exception:
            pass
    elif ext in ['.jpg', '.jpeg', '.png', '.webp']:
        pass
    else:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                raw_text = f.read()
        except Exception:
            pass

    if not raw_text.strip():
        if ext == '.docx':
            return extract_from_word(file_path)
        if ext == '.pdf':
            return extract_from_pdf(file_path)
        return []

    ai = AIService()
    prompt = f"""
    قم بتحليل النص التالي المستخرج من جدول استعمال الزمن أو قائمة إسناد الأساتذة في مدرسة.
    المطلوب:
    استخرج قائمة بجميع الأساتذة مع المواد التي يدرسونها والأقسام المسندة إليهم.
    يجب أن يكون الناتج حصرياً بصيغة JSON Array مصفوفة، حيث كل عنصر هو كائن Object يحتوي على:
    - name: اسم الأستاذ واللقب (نص)
    - subject: مادة التدريس (نص). إذا لم تكن واضحة ضع "/"
    - classes: مصفوفة نصوص بأسماء الأقسام المسندة للأستاذ (مثل ["1م1", "4م2", "3م1"])

    مثال للإخراج المطلوب:
    [
      {{"name": "محمد بن علي", "subject": "رياضيات", "classes": ["1م1", "2م2"]}},
      {{"name": "فاطمة الزهراء", "subject": "لغة عربية", "classes": ["4م1"]}}
    ]

    تنبيه: لا تكتب أي نص إضافي قبل أو بعد الـ JSON. الإخراج يجب أن يكون JSON قابل للتحليل مباشرة.

    النص للتحليل:
    {raw_text[:8000]}
    """

    response = ai.generate_response("أنت خبير في تحليل البيانات المدرسية وصياغتها بـ JSON دقيق.", prompt, rag_enabled=False)

    import json
    import re
    if response and response != "INSUFFICIENT_FUNDS_ERROR" and not response.startswith("⚠️"):
        clean_json = re.sub(r'```json\s*', '', response)
        clean_json = re.sub(r'```\s*', '', clean_json).strip()
        try:
            candidates = json.loads(clean_json)
            normalized = []
            for c in candidates:
                if isinstance(c, dict) and 'name' in c and 'classes' in c:
                    normalized.append({
                        'name': c.get('name', 'غير معروف'),
                        'subject': c.get('subject', '/'),
                        'classes': c.get('classes', [])
                    })
            if normalized:
                return normalized
        except json.JSONDecodeError as e:
            logger.error("Failed to parse AI JSON response: %s\nResponse: %s", e, clean_json[:500])

    if ext == '.docx':
        return extract_from_word(file_path)
    if ext == '.pdf':
        return extract_from_pdf(file_path)
    return []
