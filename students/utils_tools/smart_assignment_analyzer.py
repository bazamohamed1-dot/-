import pandas as pd
import re
import os
import logging
from docx import Document

logger = logging.getLogger(__name__)

# Pattern to detect any variation of class name:
# e.g., "1 م 1", "1M1", "1 AM 1", "4متوسط3"
# Broad pattern: Digit + (optional text) + Digit
CLASS_REGEX = r'\b(\d+)\s*[\u0600-\u06FFa-zA-Z]*\s*(\d+)\b'

def _detect_assignment_columns(df):
    """
    Scans first 12 rows to find header row and column indices.
    Structure: الأستاذ | المادة | الرتبة | [قسم1] | [قسم2] | [قسم3] | ...
    الأقسام المسندة = كل عمود بعد آخر عمود metadata (الرتبة/المادة/الحجم الساعي).
    Returns (name_col, subject_col, last_metadata_col, data_start_row).
    """
    name_keywords = ['الأستاذ', 'الاستاذ', 'اسم الأستاذ', 'الاسم', 'اسم', 'الأساتذة', 'الموظف']
    subject_keywords = ['المادة', 'المادة المدرسة', 'المادة الدراسية', 'التخصص']
    metadata_keywords = ['الرقم', 'الأستاذ', 'الاستاذ', 'الرتبة', 'المادة', 'الحجم الساعي', 'الحصص الأسبوعية', 'التوقيت']

    best = {'name_col': None, 'subject_col': None, 'last_metadata_col': -1, 'data_start': 0, 'score': 0}

    for row_idx in range(min(12, len(df))):
        row = df.iloc[row_idx]
        row_str = [str(c).strip() if pd.notna(c) and str(c).lower() != 'nan' else '' for c in row.values]
        name_col, subject_col = None, None
        last_meta = -1

        for col_idx, cell in enumerate(row_str):
            if not cell or len(cell) < 2:
                continue
            cell_lower = cell.lower()
            for kw in name_keywords:
                if kw in cell or kw in cell_lower:
                    name_col = col_idx
                    last_meta = max(last_meta, col_idx)
                    break
            for kw in subject_keywords:
                if kw in cell or kw in cell_lower:
                    subject_col = col_idx
                    last_meta = max(last_meta, col_idx)
                    break
            for kw in metadata_keywords:
                if kw in cell or kw in cell_lower:
                    last_meta = max(last_meta, col_idx)
                    break

        score = (1 if name_col is not None else 0) + (1 if subject_col is not None else 0)
        if score >= 1 and last_meta >= 0 and (score > best['score'] or (score == best['score'] and last_meta > best['last_metadata_col'])):
            best = {'name_col': name_col, 'subject_col': subject_col, 'last_metadata_col': last_meta,
                    'data_start': row_idx + 1, 'score': score}

    if best['score'] >= 1 and best['name_col'] is not None:
        return best['name_col'], best['subject_col'], best['last_metadata_col'], best['data_start']
    return None, None, -1, 0


def _is_likely_class_code(val):
    """Returns True if the value looks like a class code (1م4, أولى 1, 4-2, etc)."""
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return False
    s = str(val).strip()
    if not s or len(s) > 20:
        return False
    if re.match(r'^-?\d+\.?\d*$', s):
        return False
    if re.match(r'^\d+م\d+$', s.replace(' ', '')):
        return True
    if re.match(r'^\d+\s*[مM]\s*\d+$', s, re.I):
        return True
    if re.search(r'أولى|ثانية|ثالثة|رابعة|متوسط', s):
        return True
    if re.match(r'^\d+[-/]\d+$', s):
        return True
    return False


def _parse_single_cell_as_class(cell_value):
    """Treats cell as single class code or extracts from it. Returns list of class codes."""
    if cell_value is None or (isinstance(cell_value, float) and pd.isna(cell_value)):
        return []
    s = str(cell_value).strip()
    if not s:
        return []
    parsed = _parse_classes_cell(s)
    if parsed:
        return parsed
    if _is_likely_class_code(s):
        normalized = re.sub(r'\s+', '', s)
        if re.match(r'^\d+م\d+$', normalized):
            return [normalized]
        m = re.match(r'^(\d+)[-/](\d+)$', s)
        if m:
            return [f"{m.group(1)}م{m.group(2)}"]
        return [s]
    return []


def _parse_classes_cell(cell_value):
    """Extract class codes from a cell that may contain '1م4 2م4 3م4', '4م1-4م2', '1م1، 1م2', etc."""
    if cell_value is None or (isinstance(cell_value, float) and pd.isna(cell_value)):
        return []
    text = str(cell_value).strip()
    if not text:
        return []

    found = []
    raw_classes = re.findall(r'\b\d+\s*(?:M|AM|م|متوسط)\s*\d+\b', text, re.IGNORECASE)
    for rc in raw_classes:
        clean = re.sub(r'\s+', ' ', rc).strip()
        if clean:
            found.append(clean)

    word_classes = re.findall(r'\b(?:أولى|ثانية|ثالثة|رابعة|الاولى|الثانية|الثالثة|الرابعة|اولى|ثانيه|ثالثه|رابعه)\s*(?:متوسط)?\s*\d+\b', text)
    for wc in word_classes:
        clean = re.sub(r'\s+', ' ', wc).strip()
        if clean:
            found.append(clean)

    digit_digit = re.findall(r'\b(\d+)\s*[-/،,]+\s*(\d+)\b', text)
    for a, b in digit_digit:
        found.append(f"{a}م{b}")

    return list(set(found))


def normalize_class_name(raw_name):
    """
    Returns the raw name cleaned up but NOT forced to '1AM1' format unless absolutely necessary.
    The user wants to keep the original Arabic format if present.
    Example: "1 م 1" -> "1 م 1"
    """
    if not isinstance(raw_name, str):
        return None

    clean = raw_name.strip()
    # If it's just spaces, return None
    if not clean: return None

    return clean

def normalize_subject(subject):
    """
    Maps various subject names to standard database values.
    """
    if not isinstance(subject, str):
        return None

    s = subject.strip()

    subjects_map = {
        'رياضيات': 'رياضيات', 'Math': 'رياضيات',
        'فيزياء': 'فيزياء', 'Physique': 'فيزياء',
        'معلوماتية': 'إعلام آلي', 'إعلام آلي': 'إعلام آلي', 'حاسوب': 'إعلام آلي',
        'علوم الحاسوب': 'إعلام آلي', 'إعلام': 'إعلام آلي',
        'علوم طبيعية': 'علوم طبيعية', 'علوم': 'علوم طبيعية', 'Science': 'علوم طبيعية',
        'عربية': 'لغة عربية', 'Arabe': 'لغة عربية',
        'فرنسية': 'لغة فرنسية', 'Français': 'لغة فرنسية',
        'انجليزية': 'لغة إنجليزية', 'Anglais': 'لغة إنجليزية',
        'تاريخ': 'تاريخ وجغرافيا', 'Histoire': 'تاريخ وجغرافيا',
        'جغرافيا': 'تاريخ وجغرافيا', 'Géographie': 'تاريخ وجغرافيا',
        'اجتماعيات': 'تاريخ وجغرافيا',  # اجتماعيات = تاريخ وجغرافيا
        'إسلامية': 'تربية إسلامية', 'Islamique': 'تربية إسلامية',
        'مدنية': 'تربية مدنية', 'Civique': 'تربية مدنية',
        'إعلام': 'إعلام آلي', 'Informatique': 'إعلام آلي',
        'تكنولوجيا': 'تكنولوجيا', 'Technologie': 'تكنولوجيا',
        'بدنية': 'تربية بدنية', 'Sport': 'تربية بدنية',
        'موسيقى': 'تربية موسيقية', 'Musique': 'تربية موسيقية',
        'تشكيلية': 'تربية تشكيلية', 'رسم': 'تربية تشكيلية', 'Dessin': 'تربية تشكيلية',
        'رياضة': 'تربية بدنية', 'أمازيغية': 'لغة أمازيغية', 'Tamazight': 'لغة أمازيغية'
    }

    for key, val in subjects_map.items():
        if key in s:
            return val
    return None

def extract_from_excel(file_path):
    """
    Extracts assignment candidates from an Excel file (Schedule).
    First tries structure-aware parsing (detect headers: الأستاذ, المادة, الأقسام),
    then falls back to text-based extraction for unstructured files.
    """
    candidates = []
    try:
        df = pd.read_excel(file_path, header=None)

        name_col, subject_col, last_metadata_col, data_start = _detect_assignment_columns(df)
        class_cols_start = last_metadata_col + 1

        if name_col is not None and last_metadata_col >= 0:
            for idx in range(data_start, len(df)):
                row = df.iloc[idx]
                cells = row.tolist() if hasattr(row, 'tolist') else list(row)

                def safe_cell(col):
                    if col is None or col < 0 or col >= len(cells):
                        return ''
                    v = cells[col]
                    if pd.isna(v):
                        return ''
                    return str(v).strip()

                name = safe_cell(name_col)
                subject_raw = safe_cell(subject_col) if subject_col is not None else ''

                parsed_classes = []
                for col_idx in range(class_cols_start, len(cells)):
                    cell_val = cells[col_idx] if col_idx < len(cells) else None
                    parsed_classes.extend(_parse_single_cell_as_class(cell_val))

                if not parsed_classes:
                    row_text = "  ".join([str(c).strip() for c in cells if c and str(c).lower() != 'nan'])
                    parsed_classes = _extract_classes_from_text(row_text)

                subject = normalize_subject(subject_raw) if subject_raw else None
                if not subject and subject_raw:
                    subject = subject_raw.strip() if subject_raw.strip() else '/'
                elif not subject:
                    subject = '/'

                if len(name) < 2 and (parsed_classes or subject != '/'):
                    row_text = "  ".join([str(c).strip() for c in cells[:class_cols_start] if c and str(c).lower() != 'nan'])
                    for noise in ['الرقم', 'الأستاذ', 'الرتبة', 'المادة', 'الحجم الساعي', 'الأقسام']:
                        row_text = row_text.replace(noise, ' ')
                    name = re.sub(r'\s+', ' ', row_text).strip()

                if len(name) < 2:
                    continue
                if not parsed_classes and subject == '/':
                    continue

                extracted = {'name': name[:200], 'subject': subject, 'classes': list(set(parsed_classes))}
                _merge_candidate(candidates, extracted)

        if candidates:
            return candidates

        for idx, row in df.iterrows():
            row_cells = [str(x).strip() for x in row.values if str(x).lower() != 'nan' and str(x).strip() != '']
            if not row_cells:
                continue
            row_text = "  ".join(row_cells)
            extracted = _extract_candidate_from_text(row_text)
            if extracted:
                _merge_candidate(candidates, extracted)

    except Exception as e:
        logger.error(f"Excel Extraction Failed: {e}")

    return candidates


def _extract_classes_from_text(text):
    """Helper to extract class codes from joined row text."""
    found = []
    raw = re.findall(r'\b\d+\s*(?:M|AM|م|متوسط)\s*\d+\b', text, re.IGNORECASE)
    for r in raw:
        found.append(re.sub(r'\s+', ' ', r).strip())
    word = re.findall(r'\b(?:أولى|ثانية|ثالثة|رابعة|الاولى|الثانية|الثالثة|الرابعة|اولى|ثانيه|ثالثه|رابعه)\s*(?:متوسط)?\s*\d+\b', text)
    for w in word:
        found.append(re.sub(r'\s+', ' ', w).strip())
    return list(set(found))

def extract_from_pdf(file_path):
    """
    Extracts assignment candidates from a PDF document using pdfplumber.
    It reads the text layout to find Teacher Names, Subjects, and Classes.
    This saves massive AI tokens by doing local text processing.
    """
    candidates = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Try extracting tables first
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        # row is a list of strings
                        row_text = " ".join([str(cell).strip() for cell in row if cell])
                        extracted = _extract_candidate_from_text(row_text)
                        if extracted:
                            _merge_candidate(candidates, extracted)

                # Also try normal text extraction line by line in case it's not a standard table
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        extracted = _extract_candidate_from_text(line)
                        if extracted:
                            _merge_candidate(candidates, extracted)

    except Exception as e:
        logger.error(f"PDF Extraction Failed: {e}")

    return candidates

def extract_from_word(file_path):
    """
    Extracts assignment candidates from a Word document.
    Reads tables and paragraphs.
    """
    candidates = []
    try:
        doc = Document(file_path)

        # 1. Process Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        row_text.append(txt)

                line = " ".join(row_text)
                extracted = _extract_candidate_from_text(line)
                if extracted:
                    _merge_candidate(candidates, extracted)

        # 2. Process Paragraphs (if not in table)
        for para in doc.paragraphs:
            extracted = _extract_candidate_from_text(para.text)
            if extracted:
                _merge_candidate(candidates, extracted)

    except Exception as e:
        logger.error(f"Word Extraction Failed: {e}")

    return candidates

def _extract_candidate_from_text(text):
    """
    Helper to parse a single line of text for Teacher Name, Subject, Classes.
    """
    text = text.strip()
    if len(text) < 5:
        return None

    # 1. Identify Subject
    subject = normalize_subject(text)

    # 2. Identify Classes
    # Pattern: 1M1, 4AM2, 4 م 3, 1 م 1
    # We look for digit + (optional chars) + digit

    # New Regex to catch "1 م 1" specifically along with "1M1"
    # Matches: "1" then spaces then "م" or "M" or "AM" then spaces then "1"
    raw_classes = re.findall(r'\b\d+\s*(?:M|AM|م|متوسط)\s*\d+\b', text, re.IGNORECASE)

    found_classes = []
    for rc in raw_classes:
        # Keep it essentially as is, just clean spaces
        # But user wants consistency? "1 م 2" vs "1م2"
        # Let's clean extra internal spaces
        clean_c = re.sub(r'\s+', ' ', rc).strip()
        found_classes.append(clean_c)

    # Sometimes AI output or files use formats like "أولى 1", we shouldn't rely solely on "م" locally.
    # Let's expand local extraction to catch words
    word_classes = re.findall(r'\b(?:أولى|ثانية|ثالثة|رابعة|الاولى|الثانية|الثالثة|الرابعة|اولى|ثانيه|ثالثه|رابعه)\s*(?:متوسط)?\s*\d+\b', text)
    for wc in word_classes:
        clean_wc = re.sub(r'\s+', ' ', wc).strip()
        found_classes.append(clean_wc)

    # If no subject and no classes, unlikely to be a schedule row
    if not subject and not found_classes:
        return None

    # 3. Identify Teacher Name
    # Remove subject and classes from text to find the name
    clean_text = text
    if subject:
        # Remove common subject keywords
        for key in ['رياضيات', 'علوم', 'فيزياء', 'عربية', 'فرنسية', 'انجليزية', 'تاريخ', 'جغرافيا', 'إسلامية', 'مدنية', 'تكنولوجيا', 'إعلام', 'بدنية', 'موسيقى', 'تشكيلية']:
             clean_text = clean_text.replace(key, '')

    for rc in raw_classes:
        clean_text = clean_text.replace(rc, '')

    # Remove common noise words
    noise = ['الأستاذ', 'الاستاذ', 'المادة', 'القسم', 'التوقيت', 'السيد', 'السيدة', 'الآنسة', ':', '-', '/']
    for n in noise:
        clean_text = clean_text.replace(n, ' ')

    # Remove digits (hours, room numbers)
    clean_text = re.sub(r'\d+', '', clean_text)

    name = clean_text.strip()

    # Validation
    if len(name) < 3:
        return None # Name too short

    if not found_classes:
        return None # No classes assigned?

    return {
        'name': name,
        'subject': subject if subject else '/',
        'classes': list(set(found_classes))
    }

def _merge_candidate(candidates, new_cand):
    """
    Merges entries only if Name AND Subject match.
    If Name matches but Subject is different, treats as separate assignment (e.g. Teacher X - Math vs Teacher X - Physics).
    """
    for c in candidates:
        # Strict matching on Name AND Subject to support multi-subject teachers
        if c['name'] == new_cand['name']:

            # If subjects are identical, merge classes
            if c['subject'] == new_cand['subject']:
                c['classes'].extend(new_cand['classes'])
                c['classes'] = list(set(c['classes']))
                return

            # If existing has no subject ('/'), take the new subject and merge
            # This handles cases where one row has subject, another doesn't for same teacher
            if c['subject'] == '/' and new_cand['subject'] != '/':
                c['subject'] = new_cand['subject']
                c['classes'].extend(new_cand['classes'])
                c['classes'] = list(set(c['classes']))
                return

            # If new has no subject, merge into existing
            if new_cand['subject'] == '/':
                c['classes'].extend(new_cand['classes'])
                c['classes'] = list(set(c['classes']))
                return

    # If no match found (different name OR different subject), append as new
    candidates.append(new_cand)
